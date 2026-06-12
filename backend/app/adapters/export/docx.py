"""Report markdown → .docx (`python-docx`).

Walks the canonical markdown (the subset the report writer emits: `#`/`##`/
`###` headings, `-` bullets, paragraphs with `**bold**`, `*italic*` and
`[^src:id]` citation markers). Citation markers become superscript `[n]` runs
numbered by first appearance; the references section is built from the same
numbering, so the docx list always matches the inline citations. Confidence
tags like `*(confidence: emerging)*` arrive as italic spans and are preserved
as italic runs — labels survive the export (overview invariant 6).
"""

from __future__ import annotations

import io
import re

from docx import Document
from docx.text.paragraph import Paragraph

from app.adapters.export.base import Exporter, ReportContent, format_reference
from app.services import citations

_INLINE = re.compile(
    r"\[\^src:(?P<src>[0-9a-fA-F-]{8,36})\]" r"|\*\*(?P<bold>[^*]+)\*\*" r"|\*(?P<italic>[^*]+)\*"
)


def _add_inline_runs(paragraph: Paragraph, text: str, numbering: dict[str, int]) -> None:
    pos = 0
    for match in _INLINE.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        source_id = match.group("src")
        if source_id is not None:
            run = paragraph.add_run(f"[{numbering.get(source_id, '?')}]")
            run.font.superscript = True
        elif match.group("bold") is not None:
            paragraph.add_run(match.group("bold")).bold = True
        else:
            paragraph.add_run(match.group("italic")).italic = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


class DocxReportExporter(Exporter[ReportContent]):
    extension = "docx"
    media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def render(self, content: ReportContent) -> bytes:
        markdown = content.report.content_markdown or ""
        numbering = {s: i + 1 for i, s in enumerate(citations.cited_source_ids(markdown))}

        document = Document()
        for raw_line in markdown.splitlines():
            line = raw_line.rstrip()
            if not line.strip():
                continue
            if line.startswith("### "):
                _add_inline_runs(document.add_heading("", level=3), line[4:], numbering)
            elif line.startswith("## "):
                _add_inline_runs(document.add_heading("", level=2), line[3:], numbering)
            elif line.startswith("# "):
                _add_inline_runs(document.add_heading("", level=1), line[2:], numbering)
            elif line.startswith("- "):
                paragraph = document.add_paragraph(style="List Bullet")
                _add_inline_runs(paragraph, line[2:], numbering)
            else:
                _add_inline_runs(document.add_paragraph(), line, numbering)

        if numbering:
            document.add_heading("References", level=2)
            for source_id, number in sorted(numbering.items(), key=lambda kv: kv[1]):
                document.add_paragraph(
                    format_reference(number, source_id, content.sources.get(source_id))
                )

        buffer = io.BytesIO()
        document.save(buffer)
        return buffer.getvalue()
